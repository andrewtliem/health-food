import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_document():
    doc = docx.Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette
    PRIMARY = RGBColor(26, 77, 46)     # Deep Forest Green (#1A4D2E)
    SECONDARY = RGBColor(79, 111, 82)  # Muted Sage (#4F6F52)
    TEXT_DARK = RGBColor(30, 41, 59)   # Slate 800 (#1E293B)
    CODE_COLOR = RGBColor(15, 23, 42)  # Slate 900
    ACCENT_WARN = RGBColor(180, 83, 9) # Amber 700

    # Title Page / Header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("PANDUAN PRAKTIKUM E-BUSINESS & CLOUD COMPUTING\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run("Membangun Platform E-Commerce Modern dengan Cloudflare Pages, Edge Functions, Serverless SQL (D1), Payment Gateway, & Edge AI\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = SECONDARY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = meta.add_run("Studi Kasus Nyata: Eden Healthy Market - Universitas Klabat (UNKLAB)\nBuku Panduan Mahasiswa untuk Membangun Mental Model Arsitektur Modern\n")
    run_meta.font.name = "Arial"
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph("―" * 45).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = SECONDARY
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.color.rgb = TEXT_DARK
        return p

    def add_body(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.color.rgb = TEXT_DARK
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r_bold = p.add_run(bold_prefix + ": ")
        r_bold.font.name = "Arial"
        r_bold.font.size = Pt(10)
        r_bold.font.bold = True
        r_bold.font.color.rgb = TEXT_DARK

        r_text = p.add_run(text)
        r_text.font.name = "Arial"
        r_text.font.size = Pt(10)
        r_text.font.color.rgb = TEXT_DARK
        return p

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F1F5F9")  # Light slate gray
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(code_text)
        r.font.name = "Courier New"
        r.font.size = Pt(8.5)
        r.font.color.rgb = CODE_COLOR
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_callout(title_text, body_text, bg_hex="ECFDF5", border_color=PRIMARY):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, bg_hex)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r_title = p.add_run("💡 " + title_text + "\n")
        r_title.font.name = "Arial"
        r_title.font.size = Pt(9.5)
        r_title.font.bold = True
        r_title.font.color.rgb = border_color
        r_body = p.add_run(body_text)
        r_body.font.name = "Arial"
        r_body.font.size = Pt(9)
        r_body.font.color.rgb = TEXT_DARK
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # -------------------------------------------------------------
    # BAB 1
    # -------------------------------------------------------------
    add_heading_1("BAB 1: PENGANTAR & TUJUAN PEMBELAJARAN (MENGATASI COGNITIVE DEBT)")
    add_body("Di era perkembangan kecerdasan buatan (Artificial Intelligence) saat ini, mahasiswa dapat dengan sangat mudah meminta AI menghasilkan ribuan baris kode dalam hitungan detik. Namun, kemudahan ini memunculkan bahaya akademik yang sangat nyata: Cognitive Debt (Hutang Kognitif).")
    
    add_callout("Apa itu Cognitive Debt?", 
        "Cognitive Debt terjadi ketika seorang perekayasa perangkat lunak mengimplementasikan kode yang 'berjalan dengan sukses' tetapi TIDAK MEMAHAMI prinsip dasar arsitekturnya, alur aliran datanya, maupun alasan teknis di balik pemilihan teknologinya. Saat terjadi error atau kebutuhan scale-up, mereka menjadi lumpuh secara teknis karena hanya bergantung pada tebakan AI.",
        bg_hex="FEF3C7", border_color=ACCENT_WARN)

    add_body("Tujuan utama praktikum ini bukan sekadar 'membuat web jadi', melainkan melatih MENTAL MODEL yang kokoh tentang bagaimana arsitektur Cloud Computing modern bekerja. Anda akan mempelajari arsitektur Jamstack + Edge Computing yang digunakan oleh startup kelas dunia.")

    add_heading_2("Mental Model: 4 Lapisan Komputasi Modern")
    add_body("Bandingkan arsitektur tradisional (Monolith LAMP: Linux, Apache, MySQL, PHP) dengan arsitektur Edge Modern yang akan kita bangun:")

    add_bullet("1. Client Browser (Frontend Runtime)", "Dijalankan langsung di gawai pengguna (Chrome/Safari). Dibangun dengan React 18, Vite, dan Tailwind CSS. Bertanggung jawab atas rendering antarmuka pengguna, animasi keranjang, interaksi checkout, dan widget chat.")
    add_bullet("2. Serverless Edge Worker (Backend API)", "Dijalankan di 330+ kota di seluruh dunia menggunakan Cloudflare Workers via direktori /functions. Menggunakan framework Hono.js. Menggantikan backend server konvensional dengan waktu start 0 milidetik (Zero Cold Starts) melalui V8 Isolates.")
    add_bullet("3. Edge Relational Database (Cloudflare D1)", "Database SQL relasional (berbasis SQLite) yang didistribusikan langsung di jaringan edge. Kode backend Anda dapat melakukan query SQL (SELECT, INSERT, UPDATE) tanpa perlu mengelola server database MySQL/PostgreSQL sendiri.")
    add_bullet("4. Edge Artificial Intelligence (Workers AI & Gemini)", "Model AI (LLM) yang berjalan di cloud. AI bertindak sebagai Asisten Layanan Pelanggan (Customer Service) pintar yang diinjeksi dengan data katalog D1 riil secara real-time sebelum memberikan jawaban ke pelanggan.")

    # -------------------------------------------------------------
    # BAB 2
    # -------------------------------------------------------------
    add_heading_1("BAB 2: STUDI KASUS BISNIS: EDEN HEALTHY MARKET (UNKLAB)")
    add_body("Sebuah sistem teknologi tidak pernah dibangun di ruang hampa. Teknologi selalu hadir untuk menyelesaikan masalah bisnis yang konkret (Business Problem Transformation).")
    
    add_heading_2("Profil Bisnis & Masalah Nyata")
    add_body("Eden Healthy Market adalah UMKM lokal yang berlokasi di Kompleks Kampus Universitas Klabat (UNKLAB), Jl. Arnold Mononutu, Airmadidi, Minahasa Utara, Sulawesi Utara. Toko ini menyediakan makanan vegetarian, produk gandum utuh, oat milk, tempe non-GMO daun pisang, sayuran organik segar, dan camilan sehat.")
    add_body("Sebelum transformasi digital, mereka menghadapi masalah operasional yang sangat melelahkan:")
    add_bullet("Masalah 1: Tanya-Jawab Stok Manual via WhatsApp", "Setiap hari admin toko harus membalas ratusan chat manual: 'Kak, oat milk masih ada?', 'Kak, granola cokelat habis belum?'. Waktu terbuang sia-sia.")
    add_bullet("Masalah 2: Ketidaktahuan Informasi Gizi & Alergen", "Mahasiswa vegetarian baru atau penderita intoleransi laktosa kesulitan membaca komposisi produk di toko fisik.")
    add_bullet("Masalah 3: Proses Checkout Tidak Terstruktur", "Pelanggan harus transfer manual, kirim bukti transfer, dan konfirmasi manual tanpa kepastian barang langsung disiapkan.")

    add_heading_2("Solusi Digital yang Kita Bangun")
    add_bullet("Katalog Digital dengan Badge Stok Real-Time", "Sistem menampilkan label otomatis: 'In Stock' (hijau), 'Low Stock' (kuning peringatan jika stok <= 5), dan 'Out of Stock' (merah). Stok berkurang otomatis di database D1 setiap kali order berhasil.")
    add_bullet("Dual-Fulfillment Khusus Mahasiswa UNKLAB", "Pelanggan dapat memilih 'Click & Collect (Pickup 1 Jam di Toko UNKLAB)' yang gratis ongkir, atau 'Same-Day Courier Delivery' dengan gratis ongkir untuk belanja di atas Rp 150.000.")
    add_bullet("Simulasi Pembayaran Midtrans Snap", "Menghadirkan pengalaman pembayaran digital standar Indonesia: QRIS (GoPay/OVO/Dana/BCA Mobile), Virtual Account BCA/Mandiri/BNI, dan Kartu Debit/Kredit.")
    add_bullet("AI Customer Assistant 24/7", "Chatbot cerdas yang menjawab pertanyaan mahasiswa tentang stok dan panduan diet vegetarian kampus menggunakan Llama 3.2 / Gemini Flash.")

    # -------------------------------------------------------------
    # BAB 3
    # -------------------------------------------------------------
    add_heading_1("BAB 3: PERSIAPAN LINGKUNGAN & REKAYASA PROYEK")
    add_heading_2("1. Prasyarat Perangkat Lunak")
    add_body("Pastikan komputer Anda telah terpasang:")
    add_bullet("Node.js", "Versi 18 LTS atau 20 LTS (unduh dari nodejs.org). Periksa dengan perintah: node -v")
    add_bullet("Git", "Sistem kontrol versi (git -v).")
    add_bullet("Akun Cloudflare", "Daftar gratis di dash.cloudflare.com (tanpa kartu kredit).")
    add_bullet("Akun GitHub", "Untuk repositori kode dan otomatisasi CI/CD.")

    add_heading_2("2. Inisialisasi Proyek Frontend")
    add_body("Proyek ini dibangun menggunakan Vite dengan React 18 dan TypeScript:")
    add_code_block("""# 1. Buat folder proyek dan inisialisasi Vite React TypeScript
npm create vite@latest health-food -- --template react-ts

# 2. Masuk ke direktori proyek
cd health-food

# 3. Pasang library yang dibutuhkan
npm install hono @hono/cloudflare-pages lucide-react
npm install -D tailwindcss postcss autoprefixer wrangler@latest

# 4. Inisialisasi konfigurasi Tailwind CSS
npx tailwindcss init -p""")

    add_heading_2("3. Struktur Direktori Proyek yang Benar")
    add_body("Pahami perbedaan folder agar tidak salah menaruh kode:")
    add_code_block("""health-food/
├── src/                    <-- KODE CLIENT (Berjalan di BROWSER)
│   ├── components/         <-- Komponen UI (Navbar, ProductGrid, Cart, ChatWidget)
│   ├── data/               <-- Mock data lokal & konstanta kategori
│   ├── lib/api.ts          <-- Helper fetch untuk memanggil backend
│   ├── types.ts            <-- Definisi interface TypeScript (Product, Order, dll)
│   ├── App.tsx             <-- Komponen induk aplikasi
│   └── main.tsx            <-- Titik masuk DOM React
├── functions/              <-- KODE SERVERLESS BACKEND (Berjalan di CLOUDFLARE EDGE)
│   └── api/
│       └── [[route]].ts    <-- Edge Router Hono.js (/api/*)
├── schema.sql              <-- Skema tabel database Cloudflare D1
├── seed.sql                <-- Data awal produk & kategori
├── wrangler.toml           <-- Konfigurasi infrastruktur (Database D1 & Workers AI)
└── package.json            <-- Dependensi dependensi npm""")

    # -------------------------------------------------------------
    # BAB 4
    # -------------------------------------------------------------
    add_heading_1("BAB 4: MERANCANG DATABASE EDGE (CLOUDFLARE D1)")
    add_body("Cloudflare D1 adalah serverless database relasional berbasis SQLite. Berbeda dari database konvensional yang berjalan di satu server terpusat, D1 membaca query dari data center Cloudflare yang paling dekat dengan pengguna.")

    add_heading_2("1. Membuat Database D1 via Terminal")
    add_body("Jalankan perintah berikut di terminal komputer Anda:")
    add_code_block("npx wrangler d1 create eden-healthy-db")
    add_body("Terminal akan menampilkan output seperti berikut:")
    add_code_block("""✅ Successfully created DB 'eden-healthy-db'!
[[d1_databases]]
binding = "DB"
database_name = "eden-healthy-db"
database_id = "b4acc9d9-ce54-458b-a074-92ef9557f9d5" """)
    add_body("Salin blok konfigurasi di atas dan simpan ke dalam file wrangler.toml Anda!")

    add_heading_2("2. Merancang Skema Database (schema.sql)")
    add_body("Kita membuat 5 tabel utama: categories, dietary_tags, products, product_dietary_tags (relasi many-to-many), dan orders:")
    add_code_block("""CREATE TABLE categories (
  id TEXT PRIMARY KEY,
  name TEXT NOT NULL,
  slug TEXT UNIQUE NOT NULL,
  icon TEXT NOT NULL,
  description TEXT
);

CREATE TABLE dietary_tags (
  id TEXT PRIMARY KEY,
  name TEXT NOT NULL,
  slug TEXT UNIQUE NOT NULL,
  badge_color TEXT NOT NULL
);

CREATE TABLE products (
  id TEXT PRIMARY KEY,
  name TEXT NOT NULL,
  slug TEXT UNIQUE NOT NULL,
  description TEXT NOT NULL,
  category_id TEXT NOT NULL REFERENCES categories(id),
  price INTEGER NOT NULL,          -- Disimpan dalam Rupiah (misal: 65000)
  unit TEXT NOT NULL,               -- Misal: '1 kg', '350 g', '1 Liter'
  stock_quantity INTEGER NOT NULL DEFAULT 0,
  image_url TEXT NOT NULL,
  origin TEXT NOT NULL,
  ingredients TEXT NOT NULL,
  allergens TEXT,
  nutritional_highlights TEXT,
  is_featured INTEGER DEFAULT 0,
  is_bundle INTEGER DEFAULT 0,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE product_dietary_tags (
  product_id TEXT NOT NULL REFERENCES products(id) ON DELETE CASCADE,
  tag_id TEXT NOT NULL REFERENCES dietary_tags(id) ON DELETE CASCADE,
  PRIMARY KEY (product_id, tag_id)
);

CREATE TABLE orders (
  id TEXT PRIMARY KEY,
  order_number TEXT UNIQUE NOT NULL,
  customer_name TEXT NOT NULL,
  customer_phone TEXT NOT NULL,
  customer_email TEXT,
  fulfillment_type TEXT NOT NULL,   -- 'pickup' | 'delivery'
  pickup_time_slot TEXT,
  delivery_address TEXT,
  items_json TEXT NOT NULL,
  subtotal INTEGER NOT NULL,
  delivery_fee INTEGER NOT NULL DEFAULT 0,
  total_amount INTEGER NOT NULL,
  payment_method TEXT NOT NULL,     -- 'qris', 'bca_va', 'mandiri_va', 'credit_card'
  payment_status TEXT NOT NULL DEFAULT 'pending',
  order_status TEXT NOT NULL DEFAULT 'processing',
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP
);""")

    add_heading_2("3. Menjalankan Migrasi Skema & Seeding Data")
    add_callout("PENTING: Masalah Komentar SQL di Dashboard", 
        "Jangan pernah meng-copy file SQL yang diawali baris komentar (-- comment) langsung ke Console browser Cloudflare. Engine web console Cloudflare akan menganggap seluruh query adalah komentar dan menghasilkan error 'The request is malformed'. Gunakan selalu perintah CLI berikut!",
        bg_hex="FEE2E2", border_color=RGBColor(220, 38, 38))

    add_body("Jalankan kedua perintah ini secara berurutan di terminal:")
    add_code_block("""# 1. Eksekusi pembuatan tabel
npx wrangler d1 execute eden-healthy-db --remote --file=./schema.sql -y

# 2. Eksekusi pengisian 17 data produk awal
npx wrangler d1 execute eden-healthy-db --remote --file=./seed.sql -y

# 3. Uji apakah data sudah masuk
npx wrangler d1 execute eden-healthy-db --remote --command="SELECT name, price, stock_quantity FROM products LIMIT 5;" """)

    # -------------------------------------------------------------
    # BAB 5
    # -------------------------------------------------------------
    add_heading_1("BAB 5: MEMBANGUN SERVERLESS BACKEND (HONO DI EDGE)")
    add_body("Banyak pemula mengira API harus dibuat dengan Express.js di server Node.js. Di Cloudflare Edge, kita menggunakan Hono.js (bahasa Jepang yang berarti 'api/flame'). Hono memiliki bobot kurang dari 15KB dan berjalan dengan kecepatan instan.")

    add_heading_2("1. Mengapa ditaruh di functions/api/[[route]].ts?")
    add_body("Nama file [[route]].ts adalah fitur Wildcard Routing milik Cloudflare Pages. Semua HTTP Request yang diawali dengan /api/ (seperti /api/products, /api/orders, /api/chat) akan otomatis diarahkan ke file TypeScript ini dan dieksekusi sebagai Edge Worker.")

    add_heading_2("2. Anatomi Kode Backend")
    add_code_block("""import { Hono } from 'hono';
import { cors } from 'hono/cors';
import { handle } from 'hono/cloudflare-pages';

type Bindings = {
  DB?: any;          // Binding ke Cloudflare D1
  AI?: any;          // Binding ke Cloudflare Workers AI
  GEMINI_API_KEY?: string;
};

const app = new Hono<{ Bindings: Bindings }>().basePath('/api');
app.use('*', cors());

// GET /api/products: Mengambil data dengan filter kategori & stok
app.get('/products', async (c) => {
  const category = c.req.query('category');
  const inStockOnly = c.req.query('inStock') === 'true';

  if (c.env.DB) {
    let query = "SELECT * FROM products WHERE 1=1";
    const params = [];
    if (category && category !== 'all') {
      query += " AND category_id = ?";
      params.push(category);
    }
    if (inStockOnly) {
      query += " AND stock_quantity > 0";
    }
    const stmt = c.env.DB.prepare(query);
    const { results } = await stmt.bind(...params).all();
    return c.json({ success: true, count: results.length, data: results });
  }
  return c.json({ success: true, data: [] });
});

// POST /api/orders: Membuat pesanan dan memotong stok otomatis
app.post('/orders', async (c) => {
  const body = await c.req.json();
  const { customer_name, customer_phone, items, fulfillment_type } = body;

  // Kurangi stok barang yang dibeli secara atomik
  if (c.env.DB) {
    for (const item of items) {
      await c.env.DB.prepare(
        "UPDATE products SET stock_quantity = MAX(0, stock_quantity - ?) WHERE id = ?"
      ).bind(item.quantity, item.id).run();
    }
  }

  return c.json({ success: true, order_id: "ORD-" + Date.now() });
});

export const onRequest = handle(app);
export default app;""")

    # -------------------------------------------------------------
    # BAB 6
    # -------------------------------------------------------------
    add_heading_1("BAB 6: SIMULASI PAYMENT GATEWAY (MIDTRANS SNAP)")
    add_body("Dalam e-commerce Indonesia, payment gateway adalah jembatan finansial antara merchant, bank, dan dompet digital (e-wallet).")

    add_heading_2("Alur Kerja Payment Gateway")
    add_bullet("Langkah 1 (Checkout)", "Pelanggan memilih metode pembayaran di Checkout Modal (QRIS, BCA VA, Mandiri VA, atau Kartu Kredit).")
    add_bullet("Langkah 2 (Tokenization)", "Backend membuat transaksi dan menghasilkan Snap Token atau QR Code URL.")
    add_bullet("Langkah 3 (Interaksi Pelanggan)", "Midtrans Snap Modal terbuka di layar menampilkan countdown 15 menit, kode QRIS dinamis, atau nomor Virtual Account unik.")
    add_bullet("Langkah 4 (Simulasi Sukses)", "Ketika tombol 'Simulate Payment (Bayar Sekarang)' diklik, webhook simulasi menembak endpoint /api/midtrans/simulate-payment, mengubah status transaksi dari 'pending' menjadi 'settlement'.")
    add_bullet("Langkah 5 (Digital Receipt)", "Sistem langsung membuka OrderSuccessModal yang berisi struk digital dan instruksi pengambilan pesanan di counter UNKLAB.")

    # -------------------------------------------------------------
    # BAB 7
    # -------------------------------------------------------------
    add_heading_1("BAB 7: MENGINTEGRASIKAN REAL AI KE SERVERLESS EDGE")
    add_body("Banyak developer pemula meletakkan kode pemanggilan AI (Google Gemini / OpenAI) di file frontend React (src/App.tsx). Ini adalah KESALAHAN FATAL dalam rekayasa perangkat lunak.")

    add_callout("Mengapa AI Wajib di Backend (/functions)?",
        "1. Keamanan Kunci Rahasia: Jika API Key diletakkan di React, siapa pun bisa membuka browser Inspect Element (F12) ➔ Network, mencuri API Key Anda, dan menghabiskan tagihan Anda.\n" +
        "2. Grounding Data (RAG): Di backend edge, fungsi dapat membaca database D1 terlebih dahulu untuk mengetahui stok riil hari ini, lalu menyuntikkannya ke prompt AI sebelum dikirim ke LLM!",
        bg_hex="ECFDF5", border_color=PRIMARY)

    add_heading_2("Dua Mesin AI yang Kita Terapkan")
    add_bullet("Mesin 1: Cloudflare Workers AI (Gratis 100% tanpa API Key)", "Menggunakan model open-source mutakhir Meta Llama 3.2 3B Instruct (@cf/meta/llama-3.2-3b-instruct) yang berjalan langsung di GPU Cloudflare. Kuota gratis 10.000 Neurons per hari cukup untuk ~1.280 percakapan setiap hari!")
    add_bullet("Mesin 2: Google Gemini Flash (3.6 / 3.8)", "Jika Anda memiliki GEMINI_API_KEY dari Google AI Studio, sistem akan memanggil model Gemini Flash generasi terbaru.")

    add_heading_2("Implementasi Edge Handler (/api/chat)")
    add_code_block("""app.post('/chat', async (c) => {
  const { message, history } = await c.req.json();

  // 1. Ambil data produk live dari D1 untuk konteks real-time
  const { results: products } = await c.env.DB.prepare(
    "SELECT name, price, stock_quantity FROM products WHERE stock_quantity > 0 LIMIT 15"
  ).all();

  const systemPrompt = `Anda adalah Customer Service pintar Eden Healthy Market di Universitas Klabat (UNKLAB).
Jam buka toko fisik: 08:00 - 20:00 WITA.
Stok toko saat ini: ${JSON.stringify(products)}.
Bantu mahasiswa memilih makanan vegetarian dan jawab ketersediaan stok dengan ramah.`;

  // 2. Jalankan Llama 3.2 di Cloudflare Workers AI
  if (c.env.AI) {
    const aiRes = await c.env.AI.run('@cf/meta/llama-3.2-3b-instruct', {
      messages: [
        { role: 'system', content: systemPrompt },
        { role: 'user', content: message }
      ],
      max_tokens: 400
    });
    return c.json({ success: true, reply: aiRes.response, source: 'workers_ai_llama3' });
  }

  // 3. Fallback jika offline
  return c.json({ success: true, reply: "Halo dari Eden Healthy Market UNKLAB!" });
});""")

    # -------------------------------------------------------------
    # BAB 8
    # -------------------------------------------------------------
    add_heading_1("BAB 8: INFRASTRUCTURE-AS-CODE (wrangler.toml) & DEPLOYMENT")
    add_body("Alih-alih mengklik puluhan tombol di browser secara manual (yang rentan lupa dan membuat cognitive debt), kita menggunakan prinsip Infrastructure as Code (IaC) melalui file wrangler.toml.")

    add_heading_2("1. Isi File wrangler.toml")
    add_code_block("""name = "eden-healthy-market"
compatibility_date = "2024-09-23"
pages_build_output_dir = "dist"

# Binding Database D1 Relasional
[[d1_databases]]
binding = "DB"
database_name = "eden-healthy-db"
database_id = "b4acc9d9-ce54-458b-a074-92ef9557f9d5"

# Binding Cloudflare Workers AI (Llama 3.2)
[ai]
binding = "AI" """)

    add_heading_2("2. Langkah Build & Publish ke Dunia Nyata")
    add_body("Jalankan proses kompilasi dan pengiriman ke jaringan Cloudflare:")
    add_code_block("""# 1. Kompilasi frontend React menjadi aset statis di folder dist/
npm run build

# 2. Deploy ke Cloudflare Pages
npx wrangler pages deploy dist --project-name=eden-healthy-market --branch=main

# 3. Simpan dan push seluruh kode ke repositori GitHub Anda
git add .
git commit -m "feat: complete deployment on Cloudflare Pages and Workers"
git push origin main""")

    add_heading_2("3. Menghubungkan Automatic Deployment (CI/CD) dari GitHub")
    add_callout("Mental Model: Mengapa 'git push' Belum Tentu Otomatis Men-deploy?",
        "Banyak mahasiswa mengira bahwa hanya dengan menjalankan 'git push', Cloudflare akan langsung otomatis men-deploy web baru. INI KESALAHPAHAMAN UMUM!\n\n" +
        "Saat kita menggunakan perintah 'npx wrangler pages deploy', Cloudflare mencatat proyek tersebut sebagai proyek 'Direct Upload' (CLI-managed). Pada mode Direct Upload, Cloudflare TIDAK mendengarkan (listen) aktivitas GitHub Anda.\n\n" +
        "Agar setiap 'git push origin main' otomatis memicu kompilasi dan deployment baru tanpa perlu menjalankan wrangler di terminal, Anda harus mengaktifkan integrasi CI/CD.",
        bg_hex="FEF3C7", border_color=ACCENT_WARN)

    add_body("Ada 2 metode standar industri untuk mengaktifkan Automatic Deployment:")

    add_heading_3("Metode A: Cloudflare Native Git Integration (Paling Direkomendasikan)")
    add_body("Metode ini menggunakan Cloudflare Pages GitHub App resmi dari Cloudflare (Zero-Token Setup):")
    add_bullet("Langkah 1", "Buka Cloudflare Dashboard (dash.cloudflare.com) ➔ Workers & Pages.")
    add_bullet("Langkah 2", "Klik tombol 'Create application' (atau Create) ➔ pilih tab 'Pages' ➔ klik 'Connect to Git'.")
    add_bullet("Langkah 3", "Pilih akun GitHub Anda dan pilih repositori 'health-food' yang sudah Anda push tadi.")
    add_bullet("Langkah 4", "Konfigurasi Pengaturan Build: Framework preset: 'None' (atau Vite), Build command: 'npm run build', Build output directory: 'dist', Root directory: biarkan kosong.")
    add_bullet("Langkah 5", "Klik 'Save and Deploy'. Cloudflare akan langsung mengkloning repo Anda, menjalankan npm run build di server Cloudflare, dan menerbitkan situs.")
    add_body("Keuntungan Metode A: Setiap kali mahasiswa membuat perubahan dan melakukan 'git push origin main', Cloudflare akan mendeteksi commit baru secara otomatis via webhook dan melakukan re-deploy dalam ~45 detik! Selain itu, setiap Pull Request baru akan otomatis dibuatkan link 'Preview Deployment'.")

    add_heading_3("Metode B: GitHub Actions Workflow (.github/workflows/deploy.yml)")
    add_body("Jika Anda ingin proses build dijalankan di runner GitHub Actions (Linux Ubuntu virtual machine):")
    add_body("Buat file .github/workflows/deploy.yml di dalam repositori Anda:")
    add_code_block("""name: Deploy to Cloudflare Pages

on:
  push:
    branches:
      - main

jobs:
  deploy:
    runs-on: ubuntu-latest
    name: Build & Deploy
    steps:
      - uses: actions/checkout@v4
      - uses: actions/setup-node@v4
        with:
          node-version: 20
      - run: npm ci
      - run: npm run build
      - name: Deploy to Cloudflare
        uses: cloudflare/wrangler-action@v3
        with:
          apiToken: ${{ secrets.CLOUDFLARE_API_TOKEN }}
          accountId: ${{ secrets.CLOUDFLARE_ACCOUNT_ID }}
          command: pages deploy dist --project-name=eden-healthy-market --branch=main""")
    add_body("Catatan Metode B: Mahasiswa harus membuat API Token di Cloudflare Dashboard (My Profile ➔ API Tokens ➔ Create Token ➔ Edit Cloudflare Workers) dan menyimpannya di GitHub Repository ➔ Settings ➔ Secrets and variables ➔ Actions.")

    # -------------------------------------------------------------
    # BAB 9
    # -------------------------------------------------------------
    add_heading_1("BAB 9: PANDUAN TROUBLESHOOTING MANDIRI (DEBUGGING GUIDE)")
    add_body("Sebagai calon perekayasa perangkat lunak, kemampuan mendiagnosa masalah adalah pembeda utama antara pemula dan profesional.")

    add_heading_2("1. Error: 'The request is malformed' saat inject SQL")
    add_body("Penyebab: Baris pertama file SQL diawali dengan -- comment. Konsol web menganggap tidak ada query yang dikirim. Solusi: Gunakan selalu npx wrangler d1 execute --file=./schema.sql via CLI.")

    add_heading_2("2. Error 5028: Model AI Deprecated")
    add_body("Penyebab: Model AI seperti llama-3.1-8b sudah tidak aktif. Solusi: Ganti nama model di fungsi menjadi model yang aktif, misalnya @cf/meta/llama-3.2-3b-instruct.")

    add_heading_2("3. Rahasia Dashboard Secrets vs Direct Upload")
    add_body("Penyebab: Secrets yang diisi di dashboard Cloudflare hanya aktif pada Git Build. Jika Anda men-deploy via wrangler pages deploy, gunakan perintah: npx wrangler pages secret put NAMA_SECRET --project-name=nama-proyek.")

    # -------------------------------------------------------------
    # BAB 10
    # -------------------------------------------------------------
    add_heading_1("BAB 10: TUGAS MANDIRI & UJI PEMAHAMAN MAHASISWA")
    add_body("Untuk membuktikan bahwa Anda tidak mengalami Cognitive Debt, kerjakan latihan mandiri berikut:")
    add_bullet("Tugas 1 (Manipulasi Data D1)", "Tambahkan 1 produk baru (misalnya: 'Kombucha Apel Manado') ke dalam tabel products di D1 menggunakan perintah SQL INSERT via terminal. Pastikan produk langsung muncul di halaman web!")
    add_bullet("Tugas 2 (Eksperimen Prompt AI)", "Modifikasi systemPrompt di /functions/api/[[route]].ts agar AI selalu menambahkan kata penutup yang ramah khas Manado, misalnya: 'Makase banya, selamat berbelanja di UNKLAB!'. Uji hasilnya di ChatWidget.")
    add_bullet("Tugas 3 (Analisis Arsitektur)", "Gambarkan diagram alur data di buku catatan Anda: mulai dari saat pelanggan mengklik tombol 'Pay Now' di Midtrans Snap Modal, hingga pengurangan stok terjadi di tabel products Cloudflare D1.")

    add_callout("Catatan Cakupan Proyek (Project Scope Note)",
        "Implementasi ini berfokus pada fondasi teknis web-implementation (arsitektur Full-Stack Edge Computing, Serverless API, Database D1, simulasi Payment Gateway, dan integrasi Edge AI).\n\n" +
        "Solusi ini merupakan bagian inti fungsionalitas sistem dan BELUM MENCAKUP optimasi pemasaran digital tingkat lanjut seperti:\n" +
        "• SEO (Search Engine Optimization) & riset kata kunci (keywords)\n" +
        "• GEO (Generative Engine Optimization) — optimasi agar dikutip dan direferensikan oleh AI search engines (ChatGPT Search, Perplexity, Google AI Overviews)\n" +
        "• Meta Tags Mendalam & Structured Data (Open Graph, Twitter Card, Schema.org JSON-LD Microdata)\n" +
        "• Conversion Tracking & Web Analytics.",
        bg_hex="F1F5F9", border_color=PRIMARY)

    # Simpan dokumen
    filename = "PANDUAN_PRAKTIKUM_MAHASISWA.docx"
    doc.save(filename)
    print(f"File {filename} successfully generated!")

if __name__ == "__main__":
    create_document()
